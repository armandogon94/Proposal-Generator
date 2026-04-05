import os
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor

from app.config import settings

ASSETS_DIR = Path(__file__).parent.parent.parent / "assets"


def _hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _set_cell_shading(cell, color: str):
    hex_color = color.lstrip("#")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


class DOCXService:
    def __init__(self):
        self.primary_color = _hex_to_rgb(settings.primary_color)
        self.text_color = _hex_to_rgb("#1F2937")

    async def generate(self, proposal) -> str:
        os.makedirs(settings.exports_dir, exist_ok=True)

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = self.text_color

        # Heading styles
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = self.primary_color
            heading_style.font.name = "Calibri"

        # Cover page
        self._add_cover_page(doc, proposal)

        # Sections
        for section in sorted(proposal.sections, key=lambda s: s.order_index):
            doc.add_heading(section.title, level=1)
            for paragraph_text in section.content.split("\n\n"):
                if paragraph_text.strip():
                    if paragraph_text.strip().startswith("- ") or paragraph_text.strip().startswith("* "):
                        for line in paragraph_text.strip().split("\n"):
                            line = line.lstrip("- *").strip()
                            if line:
                                doc.add_paragraph(line, style="List Bullet")
                    else:
                        doc.add_paragraph(paragraph_text.strip())

        # Pricing table
        if proposal.pricing_items:
            self._add_pricing_table(doc, proposal)

        # Diagrams
        for diagram in proposal.diagrams:
            if diagram.svg_output:
                doc.add_heading(diagram.title or "Diagram", level=2)
                doc.add_paragraph("[Diagram - see PDF export for visual]")

        # Signature block
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 40)
        sig = doc.add_paragraph(f"Authorized by: {settings.company_name}")
        sig.runs[0].font.size = Pt(10)

        # Footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(f"{settings.company_url}")
        run.font.size = Pt(8)
        run.font.color.rgb = _hex_to_rgb("#6B7280")

        filepath = os.path.join(
            settings.exports_dir, f"{proposal.proposal_number}.docx"
        )
        doc.save(filepath)
        return filepath

    def _add_cover_page(self, doc: Document, proposal):
        # Logo
        logo_path = ASSETS_DIR / "images" / "logo.png"
        if logo_path.exists():
            doc.add_picture(str(logo_path), width=Inches(2.0))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        doc.add_paragraph("")
        title = doc.add_heading(proposal.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Proposal number
        num_para = doc.add_paragraph()
        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = num_para.add_run(proposal.proposal_number)
        run.font.color.rgb = _hex_to_rgb(settings.secondary_color)
        run.font.size = Pt(14)

        # Client info
        doc.add_paragraph("")
        client_para = doc.add_paragraph()
        client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if proposal.client:
            run = client_para.add_run(f"Prepared for: {proposal.client.name}")
            run.font.size = Pt(14)
            if proposal.client.company:
                doc.add_paragraph("")
                company_para = doc.add_paragraph()
                company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = company_para.add_run(proposal.client.company)
                run.font.size = Pt(12)
                run.font.color.rgb = _hex_to_rgb("#6B7280")

        doc.add_page_break()

    def _add_pricing_table(self, doc: Document, proposal):
        doc.add_heading("Pricing", level=1)

        items = [i for i in proposal.pricing_items if i.is_selected or not i.is_optional]
        table = doc.add_table(rows=1 + len(items) + 1, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ["Item", "Qty", "Unit Price", "Total"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            _set_cell_shading(cell, settings.primary_color)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)

        # Data rows
        for i, item in enumerate(items):
            row = table.rows[i + 1]
            row.cells[0].text = item.name
            row.cells[1].text = str(item.quantity)
            row.cells[2].text = f"${item.unit_price:,.2f}"
            row.cells[3].text = f"${item.total_price:,.2f}"
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

        # Total row
        total_row = table.rows[-1]
        total_cell = total_row.cells[0]
        total_cell.text = "Total"
        _set_cell_shading(total_cell, "#F3F4F6")
        for j in range(1, 3):
            _set_cell_shading(total_row.cells[j], "#F3F4F6")
        total_row.cells[3].text = f"${proposal.final_amount:,.2f}"
        _set_cell_shading(total_row.cells[3], "#F3F4F6")
        for cell in total_row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
